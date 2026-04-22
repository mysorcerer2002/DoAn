"""Helper builder — wrap python-docx với API gọn, tự áp style STU.

Dùng:
    from builder import ReportBuilder
    rb = ReportBuilder(title="XÂY DỰNG HỆ THỐNG TÍCH ĐIỂM…")
    rb.add_cover(sv="NGUYỄN VĂN A", mssv="XXXX", gvhd="ThS. TRẦN VĂN B", year="2025-2026")
    rb.add_toc()
    rb.start_chapter("Chương 1", "GIỚI THIỆU")
    rb.h2("1.1. Đặt vấn đề")
    rb.p("...")
    rb.save("bao-cao-final.docx")
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import style as S


class ReportBuilder:
    def __init__(self, title: str, subtitle: str = "LUẬN VĂN TỐT NGHIỆP"):
        self.title = title
        self.subtitle = subtitle
        self.doc = Document()
        S.apply_document_defaults(self.doc)
        self._fig_counter: dict[int, int] = {}  # chapter_idx -> count
        self._table_counter: dict[int, int] = {}
        self._chapter_idx = 0  # 0 = bìa/TOC; 1+ = chương 1..5
        # Xoá paragraph rỗng mặc định của Document() để bìa sạch
        for p in list(self.doc.paragraphs):
            p._element.getparent().remove(p._element)

    # ---------- Trang bìa ----------
    def add_cover(self, sv: str = "[HỌ TÊN SINH VIÊN]", mssv: str = "[MSSV]",
                  gvhd: str = "[HỌ TÊN GVHD]", year: str = "[NĂM HỌC]"):
        """Trang bìa — viền xanh dương, không header/footer."""
        sec = self.doc.sections[0]
        S.set_page_border_blue(sec)
        # Tắt header/footer cho section bìa
        sec.different_first_page_header_footer = True

        def center_para(text, size, bold=False, italic=False, underline=False):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(text)
            S._set_run_font(r, size=Pt(size), bold=bold, italic=italic, underline=underline)
            return p

        center_para("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ SÀI GÒN", 13, bold=True)
        center_para("KHOA CÔNG NGHỆ THÔNG TIN", 13, bold=True)
        center_para("---oOo---", 13)
        # Logo placeholder (2 dòng trắng)
        for _ in range(3):
            self.doc.add_paragraph()
        center_para(self.subtitle, 16, bold=True)
        for _ in range(2):
            self.doc.add_paragraph()
        center_para("Đề tài:", 14, italic=True)
        # Tên đề tài size 22
        center_para(self.title, 22, bold=True)
        for _ in range(3):
            self.doc.add_paragraph()
        center_para(f"Người hướng dẫn: {gvhd}", 13, bold=True)
        center_para(f"Sinh viên thực hiện: {sv}", 13, bold=True)
        center_para(f"MSSV: {mssv}", 13)
        for _ in range(4):
            self.doc.add_paragraph()
        center_para(f"TP. HỒ CHÍ MINH – NĂM {year}", 13, bold=True)

    # ---------- Mục lục ----------
    def add_toc(self):
        # Section mới cho mục lục — không border
        new_sec = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Reset border (copy sectPr nhưng xoá pgBorders nếu có)
        from docx.oxml.ns import qn
        sectPr = new_sec._sectPr
        pb = sectPr.find(qn("w:pgBorders"))
        if pb is not None:
            sectPr.remove(pb)
        # Không header/footer cho mục lục
        new_sec.different_first_page_header_footer = True
        for part in (new_sec.header, new_sec.footer,
                     new_sec.first_page_header, new_sec.first_page_footer):
            for p in part.paragraphs:
                p.text = ""

        title = self.doc.add_paragraph("MỤC LỤC")
        S.style_toc_title(title)
        toc_p = self.doc.add_paragraph()
        S.add_toc_field(toc_p)

    def add_list_of_figures(self):
        self.doc.add_page_break()
        title = self.doc.add_paragraph("MỤC LỤC CÁC HÌNH ẢNH")
        S.style_toc_title(title)
        p = self.doc.add_paragraph()
        # Field TOC cho Caption "Hình"
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        run = p.add_run()
        fld_b = OxmlElement("w:fldChar"); fld_b.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\h \\z \\c "Hình" '
        fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
        ph = OxmlElement("w:t"); ph.text = "Nhấn F9 để cập nhật."
        fld_e = OxmlElement("w:fldChar"); fld_e.set(qn("w:fldCharType"), "end")
        for el in (fld_b, instr, fld_sep, ph, fld_e):
            run._element.append(el)
        S._set_run_font(run, size=S.BODY_SIZE)

    # ---------- Chương ----------
    def start_chapter(self, label: str, name: str):
        """Bắt đầu chương mới (section break, header riêng, trang đầu không header).

        label = "Chương 1", name = "GIỚI THIỆU" (sẽ uppercase)
        """
        self._chapter_idx += 1
        new_sec = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Xoá border nếu có (chỉ bìa cần)
        from docx.oxml.ns import qn
        sectPr = new_sec._sectPr
        pb = sectPr.find(qn("w:pgBorders"))
        if pb is not None:
            sectPr.remove(pb)
        # Link-to-previous OFF để header riêng
        new_sec.header.is_linked_to_previous = False
        new_sec.footer.is_linked_to_previous = False
        new_sec.first_page_header.is_linked_to_previous = False
        new_sec.first_page_footer.is_linked_to_previous = False
        # Trang đầu chương KHÔNG header + KHÔNG footer trên (vẫn có footer chính theo mẫu)
        new_sec.different_first_page_header_footer = True
        S.set_header_different_first_page(new_sec)

        # Footer chung: tên đề tài + page number
        self._set_footer(new_sec, include_first_page=True)
        # Header từ trang 2 trong chương
        self._set_header(new_sec, label, name)

        # Tiêu đề chương — H1 canh phải IN HOA 24pt
        # Label riêng, sau xuống dòng, rồi tên chương
        p_label = self.doc.add_paragraph(label.upper())
        S.style_heading_1(p_label)
        p_name = self.doc.add_paragraph(name.upper())
        S.style_heading_1(p_name)
        # Bookmark để TOC nhận (dùng style Word Heading 1 thay cho custom)
        # python-docx cách đơn giản: gán style 'Heading 1' vào paragraph rồi override format
        p_label.style = self.doc.styles["Heading 1"]
        p_name.style = self.doc.styles["Heading 1"]
        S.style_heading_1(p_label)
        S.style_heading_1(p_name)

    def _set_header(self, section, label, name):
        """Header cho section: 'CHƯƠNG X: TÊN' italic uppercase."""
        # first_page_header -> để trống (trang đầu chương không header)
        for p in section.first_page_header.paragraphs:
            p.text = ""
        # header chính (các trang sau)
        header = section.header
        for p in header.paragraphs:
            p.text = ""
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.text = ""
        r = hp.add_run(f"{label.upper()}: {name.upper()}")
        S._set_run_font(r, size=S.HEADER_FOOTER_SIZE, italic=True, bold=True)

    def _set_footer(self, section, include_first_page=True):
        """Footer: tên đề tài canh trái + số trang canh phải."""
        def _fill(footer):
            for p in footer.paragraphs:
                p.text = ""
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = fp.add_run(f"Đề tài: {self.title.upper()}")
            S._set_run_font(r, size=S.HEADER_FOOTER_SIZE, italic=True)
            # Tab phải + page number
            tab_run = fp.add_run("\t")
            S._set_run_font(tab_run, size=S.HEADER_FOOTER_SIZE)
            S.add_page_number_field(fp)
            # Set tab stop phải cuối trang
            from docx.enum.text import WD_TAB_ALIGNMENT
            tab_stops = fp.paragraph_format.tab_stops
            page_width = section.page_width - section.left_margin - section.right_margin
            tab_stops.add_tab_stop(page_width, WD_TAB_ALIGNMENT.RIGHT)

        _fill(section.footer)
        if include_first_page:
            _fill(section.first_page_footer)

    # ---------- Headings ----------
    def h2(self, text: str):
        p = self.doc.add_paragraph(text.upper())
        p.style = self.doc.styles["Heading 2"]
        S.style_heading_2(p)
        return p

    def h3(self, text: str):
        p = self.doc.add_paragraph(text)
        p.style = self.doc.styles["Heading 3"]
        S.style_heading_3(p)
        return p

    def h4(self, text: str):
        p = self.doc.add_paragraph(text)
        # Heading 4 có ít style mặc định, dùng Normal
        S.style_heading_4(p)
        return p

    # ---------- Body ----------
    def p(self, text: str):
        p = self.doc.add_paragraph(text)
        S.style_body(p)
        return p

    def bullet(self, text: str):
        p = self.doc.add_paragraph(text, style="List Bullet")
        S._set_run_font(p.runs[0], size=S.BODY_SIZE) if p.runs else None
        p.paragraph_format.line_spacing = S.LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        return p

    def numbered(self, text: str):
        p = self.doc.add_paragraph(text, style="List Number")
        S._set_run_font(p.runs[0], size=S.BODY_SIZE) if p.runs else None
        p.paragraph_format.line_spacing = S.LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        return p

    # ---------- Hình ảnh ----------
    def figure(self, image_path: str, caption: str, width_cm: float = 14.0):
        """Chèn hình + caption 'Hình X-Y: ...'."""
        ch = self._chapter_idx
        self._fig_counter[ch] = self._fig_counter.get(ch, 0) + 1
        idx = self._fig_counter[ch]
        p_img = self.doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        try:
            run.add_picture(image_path, width=Cm(width_cm))
        except FileNotFoundError:
            run.add_text(f"[HÌNH MISSING: {image_path}]")
            S._set_run_font(run, size=S.BODY_SIZE, italic=True, color=None)
        # Caption với style Caption để TOC hình nhận
        p_cap = self.doc.add_paragraph()
        try:
            p_cap.style = self.doc.styles["Caption"]
        except KeyError:
            pass
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_run = p_cap.add_run(f"Hình {ch}-{idx}: ")
        S._set_run_font(num_run, size=Pt(12), bold=True, italic=True, underline=True)
        cap_run = p_cap.add_run(caption)
        S._set_run_font(cap_run, size=Pt(12), italic=True)
        return idx

    # ---------- Bảng ----------
    def table(self, headers: list[str], rows: list[list[str]], caption: str | None = None):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(h)
            S._set_run_font(r, size=S.BODY_SIZE, bold=True)
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri].cells[ci]
                cell.text = ""
                r = cell.paragraphs[0].add_run(str(val))
                S._set_run_font(r, size=S.BODY_SIZE)
        if caption:
            ch = self._chapter_idx
            self._table_counter[ch] = self._table_counter.get(ch, 0) + 1
            idx = self._table_counter[ch]
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_run = p.add_run(f"Bảng {ch}-{idx}: ")
            S._set_run_font(num_run, size=Pt(12), bold=True, italic=True, underline=True)
            cap_run = p.add_run(caption)
            S._set_run_font(cap_run, size=Pt(12), italic=True)
        return tbl

    def blank(self):
        self.doc.add_paragraph()

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path: str | Path):
        self.doc.save(str(path))
